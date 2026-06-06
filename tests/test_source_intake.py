import hashlib
import json
from pathlib import Path

from agent_service.runtime.cli import _extract_source
from agent_service.store import SQLiteStore
from agent_service.vault import Vault
from agent_service.workflows.source_intake import run_source_intake


def make_intake_vault(tmp_path: Path) -> Path:
    vault = tmp_path / "vault"
    (vault / "raw/inbox").mkdir(parents=True)
    (vault / "raw/sources").mkdir(parents=True)
    (vault / "raw/assets").mkdir(parents=True)
    (vault / "wiki").mkdir(parents=True)
    (vault / "AGENTS.md").write_text("# Agent 规则\n", encoding="utf-8")
    (vault / "purpose.md").write_text("# 目的\n", encoding="utf-8")
    (vault / "wiki/index.md").write_text("# 索引\n", encoding="utf-8")
    (vault / "wiki/log.md").write_text("# 日志\n", encoding="utf-8")
    return vault


def wiki_fingerprint(vault_path: Path) -> str:
    digest = hashlib.sha256()
    for path in sorted((vault_path / "wiki").rglob("*")):
        if path.is_file():
            digest.update(str(path.relative_to(vault_path)).encode())
            digest.update(path.read_bytes())
    return digest.hexdigest()


def test_markdown_source_intake_generates_canonical_source(tmp_path: Path):
    vault_path = make_intake_vault(tmp_path)
    input_path = tmp_path / "article.md"
    input_path.write_text("# 个人知识库\n\n这是一篇 Markdown 来源。", encoding="utf-8")
    wiki_before = wiki_fingerprint(vault_path)

    result = run_source_intake(Vault(vault_path), input_path)

    assert result.title == "个人知识库"
    assert result.format == "markdown"
    assert result.source_path.startswith("raw/sources/")
    assert result.asset_path.startswith("raw/assets/")
    assert (vault_path / result.source_path).exists()
    assert (vault_path / result.asset_path).exists()
    assert (vault_path / "system/source_manifest.json").exists()
    assert wiki_fingerprint(vault_path) == wiki_before


def test_docx_source_intake_extracts_text(tmp_path: Path):
    from docx import Document

    vault_path = make_intake_vault(tmp_path)
    input_path = tmp_path / "note.docx"
    document = Document()
    document.add_heading("DOCX 来源标题", level=1)
    document.add_paragraph("这是 DOCX 正文。")
    document.save(input_path)

    result = run_source_intake(Vault(vault_path), input_path)

    assert result.title == "DOCX 来源标题"
    assert result.format == "docx"
    source_text = (vault_path / result.source_path).read_text(encoding="utf-8")
    assert "这是 DOCX 正文。" in source_text


def test_runtime_extract_source_cli_outputs_structured_payload(tmp_path: Path):
    input_path = tmp_path / "record.md"
    input_path.write_text("# 记录文档\n\n这是一份需要进入 wiki 的文档。", encoding="utf-8")

    payload = _extract_source(str(input_path))

    assert payload["title"] == "记录文档"
    assert payload["source_path"].startswith("raw/sources/")
    assert payload["asset_path"].startswith("raw/assets/")
    assert "canonical_markdown" in payload
    assert "这是一份需要进入 wiki 的文档。" in payload["canonical_markdown"]
